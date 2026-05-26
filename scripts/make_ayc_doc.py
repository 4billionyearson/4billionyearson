"""Generate AYC Haulout Options Report as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.18)
section.right_margin  = Inches(1.18)

# ── Styles ────────────────────────────────────────────────────────────────────
def set_heading(para, text, level=1):
    para.text = text
    para.style = doc.styles[f'Heading {level}']
    run = para.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)

def add_heading(text, level=1):
    para = doc.add_heading('', level=level)
    set_heading(para, text, level)
    return para

def add_para(text, bold_start=None, space_before=0):
    p = doc.add_paragraph()
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p

def add_bullet(text, bold_start=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p

def add_table(headers, rows, caption=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3964')
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row):
            cells[ci].text = str(cell_text)
            if ri % 2 == 0:
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E9EEF6')
                tcPr.append(shd)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.italic = True
        cp.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()

def add_callout(heading_text, body_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f'{heading_text}  ')
    r.bold = True
    p.add_run(body_text)
    # light-blue shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D6E4F7')
    pPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('Arun Yacht Club', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

subtitle = doc.add_paragraph('Boat Haulout Facility — Options Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)

date_p = doc.add_paragraph('May 2026')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(10)
date_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Executive Summary', 1)

add_callout('Favoured Option:', 'Tractor and Hydraulic Boat Trailer System')

add_para(
    'The most practical and financially prudent solution for AYC is the acquisition of a '
    'compact tractor paired with a hydraulic adjustable marine boat trailer/cradle. This enables '
    'on-demand haulout and launch for all club yachts up to 32ft, operated by trained volunteer '
    'bosuns using the existing main slipway — with no requirement for new infrastructure or '
    'outside contractors.'
)
add_para(
    'Recommended approach: Purchase a quality used compact tractor (£6k–£18k) and a new or '
    'quality used hydraulic marine trailer/cradle (£8k–£20k), with total capital outlay in the '
    'range of £18,000–£35,000.'
)

add_table(
    ['Item', 'Cost (recommended route)'],
    [
        ['Used compact tractor (35–60hp, 4WD)', '£6,000–£18,000'],
        ['Marine hydraulic adjustable boat trailer', '£8,000–£20,000'],
        ['Slipway north-side tie-down improvements', '£1,000–£3,000'],
        ['Pressure washer (jet wash facility)', '£500–£1,500'],
        ['H&S assessment, training, insurance uplift', '£1,000–£2,000'],
        ['TOTAL', '£16,500–£44,500'],
    ],
    caption='Estimated capital costs — recommended procurement route.'
)

add_para(
    'At current crane pricing of approximately £325–£350 per boat per annual event (in + out '
    'combined), the club already absorbs £8,000–£9,000/year in effective cost. On-demand '
    'self-operated haulout with modest owner fees would recover the capital cost within '
    '4–6 years from event savings alone, without factoring in additional ad-hoc haulout income. '
    'This compares favourably with the recently approved £40k rib purchase.'
)

add_para('This option:')
for item in [
    'Uses the existing main slipway — no new infrastructure required',
    'Can be operated by trained volunteers with a recognised tractor competency certificate',
    'Has low annual running costs (£500–£1,500)',
    'Can be extended to offer jet washing, mast stepping, and hardstanding storage',
    'Preserves the majority of reserves for unfunded long-term maintenance (piles, pontoons)',
]:
    add_bullet(item)

doc.add_paragraph()
add_para(
    'Recommended immediate action: Form a small working group, obtain a slipway survey, '
    'approach insurers, and circulate sourcing enquiries to the RYA club network and UK marine '
    'trade press.',
    bold_start='Recommended immediate action:'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Background and Current Situation', 1)

add_para(
    'Arun Yacht Club (AYC) is a CASC-registered members\' club at Littlehampton, operated on a '
    'volunteer basis. The club\'s facilities include:'
)
for item in [
    'Main slipway (steep gradient, 24v winch, next to clubhouse) — currently used for Limbo '
     'and lifting-keel yachts; available for short-term owner bookings for maintenance. '
     'North-side mooring is improvised (lines to balcony and piles).',
    'Dinghy pen slipway (shallow gradient, 24v winch) — used for sailing dinghies; not '
     'currently deep enough for yachts, though the club owns a dredger.',
    'Car park / hardstanding — northern and southern areas used for craned-out yacht winter '
     'storage.',
]:
    add_bullet(item)

add_para(
    'Current crane arrangement: An annual crane-out in October and crane-in in March, using '
    'hired 60T and 100T cranes for approximately 20–25 yachts at £150–£175 per event per boat. '
    'Owners frequently request a summer haulout option for light maintenance and antifouling.',
    bold_start='Current crane arrangement:'
)

add_para(
    'Club profile: 80+ yacht berths; approx. 5–8 yachts race weekly; 60+ use moorings for '
    'storage only; up to 10 yachts in very poor condition occupying berths. '
    '£250k reserves; up to £40k/year surplus; significant unfunded long-term maintenance needs. '
    'Mooring fees subsidise all club overheads.',
    bold_start='Club profile:'
)

add_para(
    'Goal: An independent, on-demand facility for haulout and launch of yachts up to 32ft, '
    'operated by club maintenance/bosun volunteers, at a cost broadly equivalent to or cheaper '
    'than current crane hire.',
    bold_start='Goal:'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# OPTION 1
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Option 1: Tractor and Hydraulic Boat Trailer  (Recommended)', 1)

add_heading('Description', 2)
add_para(
    'A compact tractor (35–60hp, 4WD) paired with a hydraulic adjustable marine boat '
    'trailer/cradle tows yachts up and down the main slipway. The trailer is submerged at '
    'the base of the slip, the boat is floated over it, secured in the cradle, and hauled up. '
    'Yachts are parked on the hardstanding for maintenance or stored for longer periods.'
)
add_para(
    'This is the most common and proven solution for small-to-medium UK yacht clubs. It is '
    'used extensively on the south coast (e.g. Dell Quay SC, Emsworth SC, Hayling Island SC) '
    'and is well-suited to AYC\'s scale, budget, and volunteer operating model.'
)

add_heading('Infrastructure Requirements', 2)
for item in [
    'Main slipway: Already adequate for Limbo yachts. A structural survey is recommended '
     '(est. £500–£1,500) to confirm load-bearing capacity and surface condition.',
    'North-side tie-down points: Replace improvised arrangement with proper mooring cleats '
     'or posts (est. £1,000–£3,000).',
    'Hardstanding: Existing northern and southern car park areas already used. Painted bays '
     'and tie-down anchors may be beneficial.',
    'Trailer/tractor storage: Covered lean-to to protect equipment (est. £2,000–£5,000 if new).',
]:
    add_bullet(item)

add_heading('Equipment and Sourcing', 2)

add_para('Tractor (35–60hp, 4WD, compact utility):', bold_start='Tractor (35–60hp, 4WD, compact utility):')
add_table(
    ['Route', 'Specification', 'Cost', 'Sources'],
    [
        ['New', 'Kubota M/B series, John Deere 3000, New Holland Boomer', '£18,000–£45,000', 'Agricultural dealers, manufacturer dealers'],
        ['Used', '5–15 years old, same spec', '£6,000–£18,000', 'eBay, Machinery Trader, BIGGA, farm dispersal sales'],
    ]
)

add_para('Marine Hydraulic Adjustable Boat Trailer/Cradle:', bold_start='Marine Hydraulic Adjustable Boat Trailer/Cradle:')
add_para(
    'The trailer must be fully galvanised for saltwater use, with hydraulically adjustable '
    'padded support arms to accommodate bilge keelers, fin keelers, and lifting-keel boats.'
)
add_table(
    ['Route', 'Type', 'Cost', 'Sources'],
    [
        ['New custom', 'UK marine engineering fabricator', '£12,000–£25,000', 'Strickland Marine, W&J Marine Trailers, local fabricators'],
        ['New standard', 'European marine trailer (Balbi, Cargotec)', '£8,000–£18,000', 'UK marine equipment importers'],
        ['Used', 'Club or marina surplus', '£3,000–£12,000', 'RYA network, Boats.net, Apollo Duck, marina disposals'],
    ]
)

add_para('Additional items:')
add_table(
    ['Item', 'Estimated Cost'],
    [
        ['Pressure washer (petrol/3-phase, 200+ bar)', '£500–£1,500'],
        ['Tie-down straps, safety chains, chocks', '£500–£1,000'],
        ['Cradle pads and bilge supports (bespoke)', '£500–£1,500'],
        ['Signage, safety barriers', '£200–£500'],
    ]
)

add_heading('Operation', 2)
for item in [
    'Operators: Minimum two trained volunteers per lift — one driving the tractor, one supervising/guiding.',
    'Licensing: Tractor use on private club land does not require a public road licence. '
     'All operators should complete a recognised competency programme (e.g. Lantra Awards '
     'Tractor and Machinery Operation, approx. £150–£300 per person).',
    'Procedure: A written club Standard Operating Procedure (SOP) and risk assessment must '
     'be produced before any lifts are undertaken.',
    'Capacity: Realistically 3–6 lifts per day with two trained volunteers.',
]:
    add_bullet(item)

add_heading('Maintenance', 2)
add_table(
    ['Item', 'Estimated Annual Cost'],
    [
        ['Annual tractor service and hydraulic fluid check', '£300–£600'],
        ['Trailer inspection, greasing, galvanic protection', '£200–£500'],
        ['Total annual running cost', '£500–£1,500'],
    ]
)

add_heading('Insurance', 2)
for item in [
    'Notify current public liability insurer of the new mechanical haulout activity before any operation.',
    'Confirm cover for tractor use and mechanical haulout; check employer\'s liability if any paid staff involved.',
    'Introduce a formal owner consent/indemnity form for all lifts (standard practice at most clubs).',
    'Anticipated premium uplift: £200–£600/year.',
]:
    add_bullet(item)

add_heading('Financial Structure', 2)
add_para('Suggested fee schedule for individual haulouts:')
add_table(
    ['Service', 'Suggested Fee'],
    [
        ['Single haul (in or out)', '£80–£120'],
        ['Combined in + out (e.g. week antifouling)', '£150–£200'],
        ['Hardstanding storage (per day)', '£8–£15'],
        ['Hardstanding storage (per month, seasonal)', '£100–£200'],
        ['Jet wash (boat bottom)', '£25–£50'],
        ['Mast stepping / unstepping', '£30–£60'],
    ]
)

add_para(
    'Break-even: The club currently absorbs £8,000–£9,000/year in crane hire costs for '
    '20–25 boats (2 events × £150–£175/boat). Self-operating eliminates this cost. At a '
    'capital cost of £25,000 (mid-range), the facility pays for itself in approximately '
    '3–4 years from event savings alone. Individual on-demand haulouts at £150–£200 per '
    'boat generate additional income.',
    bold_start='Break-even:'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# OPTION 2
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Option 2: Marine Travel Lift (Self-Propelled Boat Hoist)', 1)

add_heading('Description', 2)
add_para(
    'A self-propelled straddle hoist ("travel lift") is the professional standard used by '
    'marinas and larger yacht clubs. It straddles the vessel on two parallel slings, lifts it '
    'clear of the water, and drives it to a hardstanding area. It operates from a dedicated '
    'concrete haulout bay or channel.'
)

add_heading('Infrastructure Requirements', 2)
add_para(
    'This is the significant constraint at AYC. A travel lift requires either:'
)
for item in [
    'A dedicated concrete haulout bay/channel dredged to sufficient depth (1.5m+ LAT), '
     'with reinforced concrete quay walls. Estimated civil engineering cost: £40,000–£120,000, '
     'plus Environment Agency / MHCLG permits, Harbour Commissioner consent, and potentially '
     'planning permission.',
    'Significant deepening and widening of the dinghy pen slipway using the club\'s existing '
     'dredger. The club would likely need Environment Agency consent and Harbour Commissioner '
     'approval. Timeline: 12–24 months for approvals, plus works costs.',
]:
    add_bullet(item)

add_heading('Equipment and Sourcing', 2)
add_para('For 32ft yachts up to approximately 10–12 tonnes, a 15–20T capacity travel lift is appropriate.')
add_table(
    ['Route', 'Specification', 'Cost'],
    [
        ['New 15T–20T', 'Marine Travelift (USA), Wise (Italy), Ascom, Cimolai', '£90,000–£180,000 + VAT'],
        ['Used 15T–20T', 'UK marina surplus, European brokerage', '£20,000–£65,000'],
    ]
)
add_para('Sources for used lifts: Boats.net, YachtWorld, specialist marine equipment brokers.')
add_para(
    'Total estimated cost including infrastructure: £60,000–£300,000+',
    bold_start='Total estimated cost including infrastructure:'
)

add_heading('Assessment', 2)
add_para(
    'A travel lift is the aspirational long-term solution and offers the greatest operational '
    'flexibility, including the potential to service boats from neighbouring clubs (additional '
    'revenue). However, the combination of civil engineering costs, permitting requirements, '
    'and capital outlay makes this disproportionate to AYC\'s current reserves and unfunded '
    'maintenance obligations. Recommended as a long-term Phase 2 investment once Option 1 is '
    'established and the club\'s financial position is stronger.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# OPTION 3
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Option 3: Enhanced Seasonal Crane Hire Programme', 1)

add_heading('Description', 2)
add_para(
    'Retain the current contracted crane hire arrangement but expand from one annual event to '
    'two or three events per year — adding a summer crane-out (June/July) for antifouling and '
    'light maintenance, with more structured booking and fee recovery.'
)

add_heading('How it Works', 2)
for item in [
    'Negotiate a multi-year framework contract with the current crane hire contractor.',
    'Add a June/July event (approx. 15–25 boats) for antifouling.',
    'Introduce an optional express half-day lift for boats wanting 2–3 days out.',
    'Improve booking and pre-payment systems.',
]:
    add_bullet(item)

add_heading('Costs', 2)
add_table(
    ['Item', 'Cost'],
    [
        ['Additional summer crane event (gross)', '£3,000–£6,000'],
        ['Owner cost per event', '£150–£250/boat'],
        ['Capital equipment', '£0'],
    ]
)

add_heading('Limitations', 2)
for item in [
    'Does not meet the stated goal of on-demand access for owners.',
    'Continues reliance on external contractor availability, scheduling, and pricing.',
    'No club asset created; no long-term cost reduction.',
    'Cannot accommodate urgent or ad-hoc individual requests.',
]:
    add_bullet(item)

add_heading('Assessment', 2)
add_para(
    'Lowest risk, lowest cost option but provides the least capability improvement. Best used '
    'as an interim measure for the current season while Option 1 is being planned and sourced, '
    'rather than as a permanent solution. Can be implemented immediately with a single call to '
    'the crane contractor.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# RISKS AND OPPORTUNITIES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Risks and Opportunities', 1)

add_heading('Risks', 2)
add_table(
    ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ['H&S incident during haul or launch', 'Medium', 'High', 'Written SOP and risk assessment; trained operators only; spotter mandatory; owner indemnity form'],
        ['Volunteer unavailability for lifts', 'Medium', 'Medium', 'Build a roster of 6+ trained volunteers; consider occasional paid cover for busy periods'],
        ['Insurance gap or claim', 'Medium', 'High', 'Notify insurers before any operation; obtain written confirmation of cover'],
        ['Trailer unsuitable for specific keel types', 'Low–Medium', 'Medium', 'Survey representative club keel types before purchase; request builder advice'],
        ['Slipway structural deficiency', 'Low', 'Medium', 'Commission survey before committing to procurement'],
        ['Planning or environmental consent needed', 'Low–Medium', 'High', 'Check with WSCC and Environment Agency before any permanent structure or dredging'],
        ['CASC compliance risk from fee income', 'Low', 'Medium', 'Seek written HMRC/CASC adviser confirmation that haulout fee income is within CASC rules'],
    ],
    caption='Risk register for Option 1 (Tractor/Trailer).'
)

add_heading('Opportunities', 2)
add_table(
    ['Opportunity', 'Notes'],
    [
        ['Additional income from non-members', 'Offer haulout services to neighbouring clubs, visiting boats, and private owners on the Arun'],
        ['Mast work and stepping facility', 'A simple A-frame gin pole (£1,000–£3,000) positioned on the hardstanding enables mast stepping/unstepping at minimal additional cost'],
        ['Jet washing facility', 'Low cost, popular, and simple to operate; dedicated pressure washer on hardstanding'],
        ['Dinghy pen rationalisation', 'Clear rarely-used dinghy storage; offer space as additional yacht hardstanding or dinghy storage for paying clubs'],
        ['Member retention and attraction', 'On-demand haulout is a significant differentiator for yacht owners considering berths'],
        ['Increased active sailing / CASC compliance', 'Easier maintenance access could encourage more boats to become seaworthy and participate in sailing, helping CASC eligibility'],
        ['Addresses poor-condition berths', 'Structured haulout access, combined with a seaworthiness policy, gives the committee a practical tool to address 10+ unusable yachts'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# OPTION COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Option Comparison', 1)

add_table(
    ['Criterion', 'Option 1: Tractor/Trailer', 'Option 2: Travel Lift', 'Option 3: Enhanced Crane Hire'],
    [
        ['Capital cost',                  '£16,000–£45,000',   '£60,000–£300,000+',     '£0'],
        ['Annual running cost',           '£500–£1,500',       '£2,000–£5,000',          '£0 (passed to owners)'],
        ['On-demand access',              'Yes',               'Yes',                    'No'],
        ['Handles all keel types to 32ft','Yes (correct cradle)','Yes',                  'Yes'],
        ['Operable by volunteers',        'Yes (with training)','Yes (with training)',   'No — external contractor'],
        ['Uses existing infrastructure',  'Mostly',            'No — new works required','Yes'],
        ['Infrastructure lead time',      '2–4 months',        '12–36 months',           'Immediate'],
        ['Payback (event savings)',        '3–6 years',         '15–30+ years',           'N/A'],
        ['Strain on reserves',            'Low–Medium',        'High',                   'None'],
        ['Long-term asset created',       'Yes',               'Yes',                    'No'],
        ['Overall recommendation',        'RECOMMENDED',       'Future Phase 2',         'Interim only'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# RECOMMENDED NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Recommended Next Steps', 1)

steps = [
    ('Immediately',     'Add a summer crane event for this season (Option 3 as a bridge) — one call to the crane contractor.'),
    ('Within 1 month',  'Appoint a working group of 2–3 members to take forward Option 1 (ideally including someone with farming/plant machinery experience).'),
    ('Within 2 months', 'Commission a structural survey of the main slipway. Contact club insurers. Obtain written CASC/HMRC confirmation on fee income.'),
    ('Within 3 months', 'Circulate sourcing enquiries to the RYA club network, Boats.net, and local agricultural dealers. Draft a club haulout SOP and risk assessment. Enrol 4–6 volunteers for Lantra tractor competency training.'),
    ('Within 6 months', 'Take first haulout with the new equipment — target before October crane-in to demonstrate capability.'),
]
for label, text in steps:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

doc.add_paragraph()

footer_p = doc.add_paragraph(
    'Report prepared May 2026. Costs are estimates based on current UK market rates and should '
    'be verified by quotation before commitment.'
)
footer_p.runs[0].font.size = Pt(9)
footer_p.runs[0].font.italic = True
footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/chrisdunn/Documents/My Apps/vs code/4billionyearson/public/AYC Haulout Options Report.docx'
doc.save(out)
print(f'Saved: {out}')
